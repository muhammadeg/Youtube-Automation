import os
from logging import Logger
from tkinter import messagebox
import random
from docx import Document
from datetime import datetime
import requests
import google.generativeai as genai
from moviepy import ImageClip, concatenate_videoclips, AudioFileClip, VideoFileClip, TextClip, CompositeVideoClip
from moviepy.video.fx.Resize import Resize
from PIL import Image, ImageOps
import tempfile
from config import PEXELS_API_KEY, CAPTIONS_FOLDER, PIXABAY_API_KEY, VISUAL_FOLDER, DEFAULT_PROMPT, DOCX_FOLDER, CROPPED_FOLDER, AUDIO_FOLDER, ELEVENLABS_API_KEY, ELEVENLABS_URL, CHUNK_SIZE
from arrays import VIDEO_KEYWORDS, IMAGE_KEYWORDS, VIDEO_KEYWORDS_FUN, VIDEO_KEYWORDS_FUN_FACTS
import pysrt
import re
from moviepy.video.io.ffmpeg_tools import ffmpeg_extract_subclip
import subprocess

timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

# Function to update the global timestamp
def update_timestamp():
    global timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    pass

# Function to generate the script and captions accurately based on audio timing
def generate_script(voice_required=False):
    try:
        # Initialize the chat model
        model = genai.GenerativeModel('gemini-1.5-flash')
        chat = model.start_chat(history=[])

        # Send the default prompt and get the response
        response = chat.send_message(DEFAULT_PROMPT)

        # Extract the dynamic title (first 4 words)
        script_text = response.text.strip()
        dynamic_title = " ".join(script_text.split()[:4])  # First 4 words as title

        # Remove the first 4 words from the script text for the body of the document
        script_body = " ".join(script_text.split()[4:])  # The rest of the script text

        # Create a Word document
        doc = Document()
        doc.add_heading(dynamic_title, 0)  # Use dynamic title as heading
        doc.add_paragraph(script_body)  # Add the rest of the script as body

        docx_filename = os.path.join(DOCX_FOLDER, f"File_1_{timestamp}.docx")
        audio_filename = os.path.join(AUDIO_FOLDER, f"File_1_{timestamp}.mp3")
        captions_filename = os.path.join(CAPTIONS_FOLDER, f"File_1_{timestamp}.srt")

        # Save the script to a docx file
        doc.save(docx_filename)
        print(f"Script generated and saved as {docx_filename}")

        # Generate voice if required
        if voice_required:
            generate_voice(docx_filename, audio_filename)

        # Generate captions based on audio timing
        if voice_required and os.path.exists(audio_filename):
            audio_clip = AudioFileClip(audio_filename)
            audio_duration = audio_clip.duration

            # Split script into smaller phrases
            script_phrases = re.split(r'[.?!,;]\s+', script_body.strip())  # Split on punctuation
            total_chars = sum(len(phrase) for phrase in script_phrases)

            # Create captions with weighted timing
            with open(captions_filename, 'w', encoding='utf-8') as captions_file:
                start_time = 0.0
                for idx, phrase in enumerate(script_phrases):
                    # Weight duration based on phrase length
                    char_weight = len(phrase) / total_chars
                    phrase_duration = audio_duration * char_weight
                    end_time = start_time + phrase_duration

                    # Write caption in SRT format
                    captions_file.write(f"{idx + 1}\n")
                    captions_file.write(f"{format_time(start_time)} --> {format_time(end_time)}\n")
                    captions_file.write(phrase.strip() + "\n\n")

                    start_time = end_time

            print(f"Segmented captions generated and saved as {captions_filename}")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")
    pass

# Helper function to format time in SRT format
def format_time(seconds):
    mins, secs = divmod(seconds, 60)
    millis = int((seconds - int(seconds)) * 1000)
    return f"{int(mins):02}:{int(secs):02},{millis:03}"

# Function to read the docx file and return its text
def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Function to generate the voice-over from the script
def generate_voice(docx_file, audio_file):
    try:
        # Read text from the docx file
        text_from_docx = read_docx(docx_file)

        # Prepare data for the Eleven Labs API
        data = {
            "text": text_from_docx,  # Script text from the docx file
            "model_id": "eleven_multilingual_v2",  # Multilingual voice model
            "voice_id": "ogFFqOPQ9U0REMVGJdb2",  # Specific voice ID
            "voice_settings": {
                "stability": 0.5,
                "similarity_boost": 0.9
            }
        }

        headers = {
            "Accept": "audio/mpeg",
            "Content-Type": "application/json",
            "xi-api-key": ELEVENLABS_API_KEY
        }

        # Send request to the Eleven Labs API
        response = requests.post(ELEVENLABS_URL, json=data, headers=headers)

        # Check if the response is successful
        if response.status_code == 200:
            # Save the audio output to a file
            with open(audio_file, 'wb') as f:
                for chunk in response.iter_content(chunk_size=CHUNK_SIZE):
                    if chunk:
                        f.write(chunk)

            print(f"Voice generated and saved as {audio_file}")
        else:
            raise Exception(f"Failed to generate voice: {response.status_code} - {response.text}")

    except Exception as e:
        raise Exception(f"Error generating voice: {e}")


def download_images_from_pexels(query, num_images, folder):
    headers = {
        "Authorization": PEXELS_API_KEY
    }

    # Construct the API request URL
    url = f"https://api.pexels.com/v1/search?query={query}&per_page={num_images}"
    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        data = response.json()
        os.makedirs(folder, exist_ok=True)

        image_count = 0
        for i, photo in enumerate(data['photos']):
            if image_count >= num_images:
                break
            img_url = photo['src']['original']
            try:
                # Fetch the image data
                img_data = requests.get(img_url).content

                # Create a unique filename based on the timestamp and image count
                img_filename = f"{folder}/File_1_{timestamp}_{i + 1:03d}.jpg"
                with open(img_filename, 'wb') as handler:
                    handler.write(img_data)

                image_count += 1
                print(f"Image {i + 1} saved as {img_filename}.")
            except Exception as e:
                print(f"Failed to download image {i + 1}: {e}")
    else:
        print(f"Error fetching data from Pexels API: {response.status_code}")


def download_images():
    # Select a random keyword from the list
    keyword = random.choice(IMAGE_KEYWORDS)

    # Randomize the number of images (within a range, e.g., 3 to 7)
    num_images = random.randint(3, 7)

    visual_folder = f"visual"

    # Download images using the random keyword and number
    download_images_from_pexels(keyword, num_images, visual_folder)
    print(f"Downloaded {num_images} images using keyword: {keyword}")

# Function to resize images and create video from images and audio
def resize_image(image_file, target_resolution=(1920, 1080), temp_folder=None):
    """Resize the image to the target resolution while preserving aspect ratio."""
    img = Image.open(image_file)
    resized_img = ImageOps.pad(img, target_resolution, method=Image.Resampling.LANCZOS, color=(0, 0, 0))

    # Save to a temporary file to avoid overwriting the original
    if temp_folder:
        temp_file = os.path.join(temp_folder, os.path.basename(image_file))
    else:
        temp_file = f"temp_{os.path.basename(image_file)}"

    resized_img.save(temp_file, format="PNG")  # Save as PNG to retain quality
    return temp_file


def sanitize_caption(caption):
    return re.sub(r"[^\w\s.,!?-]", "", caption)  # Remove invalid characters

def sanitize_filename(filename, max_length=100):
    """
    Sanitizes the filename by removing invalid characters and limiting its length.
    """
    sanitized = "".join(c if c.isalnum() or c in " _-" else "" for c in filename)
    return sanitized[:max_length]
def split_long_caption(caption, max_length=50):
    if not caption:
        return ""
    words = caption.split()
    lines = []
    current_line = []

    for word in words:
        if len(" ".join(current_line)) + len(word) + 1 <= max_length:
            current_line.append(word)
        else:
            lines.append(" ".join(current_line))
            current_line = [word]
    if current_line:
        lines.append(" ".join(current_line))
    return "\n".join(lines)


# Helper function to check if a video file matches the desired criteria
def is_valid_video(video_file):
    video_width = video_file['width']
    video_height = video_file['height']
    return (
        video_width >= 1080
        and video_height >= 1920
        and video_width / video_height == 9 / 16
    )

# Helper function to check if a video file matches the desired criteria
def is_valid_video_pixabay(video_file):
    video_width = video_file['width']
    video_height = video_file['height']
    aspect_ratio = video_width / video_height
    return (
        video_width >= 1080
        and video_height >= 1920
        and abs(aspect_ratio - 9 / 16) < 0.02  # Allow slight deviations in aspect ratio
    )

def crop_video_to_9_16(input_file, output_file):
    video = None  # Initialize video to avoid uninitialized variable error
    try:
        # Check if the input file exists and is not empty
        if not os.path.exists(input_file) or os.path.getsize(input_file) == 0:
            raise FileNotFoundError(f"Input file {input_file} is missing or empty.")

        # Load video file
        video = VideoFileClip(input_file)

        # Ensure video has valid FPS and duration
        if not hasattr(video, 'fps') or video.fps <= 0:
            raise ValueError(f"Invalid FPS in video {input_file}.")
        if video.duration <= 0:
            raise ValueError(f"Invalid duration in video {input_file}.")

        # Limit video duration to 5 seconds
        trimmed_video = video.subclip(0, min(5, video.duration))

        # Get original width and height
        original_width, original_height = trimmed_video.size

        # Calculate crop dimensions for 9:16 aspect ratio
        target_aspect_ratio = 9 / 16
        if original_width / original_height > target_aspect_ratio:
            # Crop width for landscape video
            new_width = int(original_height * target_aspect_ratio)
            x1 = (original_width - new_width) // 2
            x2 = x1 + new_width
            cropped_video = trimmed_video.crop(x1=x1, x2=x2)
        else:
            # Crop height for portrait or square video
            new_height = int(original_width / target_aspect_ratio)
            y1 = (original_height - new_height) // 2
            y2 = y1 + new_height
            cropped_video = trimmed_video.crop(y1=y1, y2=y2)

        # Save the cropped video
        cropped_video.write_videofile(
            output_file,
            codec="libx264",
            audio_codec="aac",
            fps=min(video.fps, 30),  # Limit FPS to 30 for compatibility
            bitrate="8000k",
            preset="medium",
            ffmpeg_params=["-pix_fmt", "yuv420p"],  # Ensure compatibility
        )
        print(f"Successfully cropped and saved video: {output_file}")

    except FileNotFoundError as fnf_error:
        print(f"File error: {fnf_error}")
    except ValueError as val_error:
        print(f"Value error: {val_error}")
    except Exception as e:
        print(f"Error cropping video {input_file}: {e}")
    finally:
        if video:
            video.close()  # Ensure video resources are released


def download_and_crop_videos_from_pixabay():
    try:
        # Select random keywords and number of videos
        selected_keywords = random.sample(VIDEO_KEYWORDS_FUN_FACTS, 5)
        num_videos = random.randint(3, 4)  # Download between 3 and 4 videos per keyword

        for keyword in selected_keywords:
            print(f"Searching for videos for keyword: {keyword}")
            downloaded_count = 0
            page = 1  # Start from the first page

            while downloaded_count < num_videos:
                url = (
                    f'https://pixabay.com/api/videos/'
                    f'?key={PIXABAY_API_KEY}&q={keyword.replace(" ", "+")}&per_page=10&page={page}'
                )
                response = requests.get(url)
                response.raise_for_status()

                videos = response.json().get('hits', [])
                if not videos:
                    break  # Exit if no results

                # Create folders if they don't exist
                os.makedirs(VISUAL_FOLDER, exist_ok=True)
                os.makedirs(CROPPED_FOLDER, exist_ok=True)

                for video in videos:
                    video_files = video.get('videos', {})
                    highest_resolution = max(video_files.values(), key=lambda x: x['width'])

                    video_id = video['id']
                    video_title = video['tags']
                    sanitized_title = sanitize_filename(video_title)
                    video_filename_base = f"{timestamp}_{keyword}_{sanitized_title}"
                    video_filename = os.path.join(VISUAL_FOLDER, f"{video_filename_base}.mp4")
                    cropped_filename = os.path.join(CROPPED_FOLDER, f"{video_filename_base}_cropped.mp4")

                    if os.path.exists(cropped_filename):
                        continue  # Skip if already processed

                    video_url = highest_resolution['url']

                    # Download and trim video
                    try:
                        for attempt in range(10):  # Retry mechanism
                            try:
                                ffmpeg_command = [
                                    "ffmpeg", "-ss", "0", "-i", video_url, "-t", "5",
                                    "-vf", (
                                        "scale=1080:1920:force_original_aspect_ratio=decrease,"
                                        "pad=1080:1920:(ow-iw)/2:(oh-ih)/2:black"
                                    ),
                                    "-c:v", "libx264", "-preset", "fast", "-crf", "23", "-pix_fmt", "yuv420p",
                                    "-c:a", "aac", "-b:a", "128k", "-y", video_filename
                                ]
                                subprocess.run(ffmpeg_command, check=True)
                                break  # Success, exit retry loop
                            except subprocess.CalledProcessError as e:
                                print(f"FFmpeg failed (attempt {attempt + 1}/3): {e}")
                        else:
                            print(f"Failed to download video {video_id} after 3 attempts.")
                            continue
                    except Exception as e:
                        print(f"Failed to download video {video_id}: {e}")
                        continue

                    # Process downloaded video
                    crop_video_to_9_16(video_filename, cropped_filename)
                    downloaded_count += 1

                    if downloaded_count >= num_videos:
                        break

                page += 1  # Next page

            if downloaded_count == 0:
                print(f"No valid videos found for keyword: {keyword}.")
            else:
                print(f"Downloaded and cropped {downloaded_count} videos using keyword: {keyword} and timestamp: {timestamp}")

    except Exception as e:
        print(f"Error downloading videos from Pixabay: {e}")

def combine_pixabay(audio_file, output_file, visual_folder="cropped", current_timestamp=None):
    if not os.path.exists(audio_file):
        print(f"Error: Audio file {audio_file} not found.")
        return

    if not current_timestamp:
        print("Error: No timestamp provided for filtering videos.")
        return

    # Get video files from the visual folder matching the current timestamp
    video_files = [
        os.path.join(visual_folder, f) for f in os.listdir(visual_folder)
        if f.endswith('.mp4') and f.startswith(f"{current_timestamp}_")
    ]

    if not video_files:
        print(f"No videos found in folder: {visual_folder} for timestamp: {current_timestamp}")
        return

    # Load the audio clip to get its duration
    audio_clip = AudioFileClip(audio_file)
    max_duration = audio_clip.duration

    # Create clips from each video, taking only the first 3 seconds, and limit total duration
    clips = []
    total_duration = 0

    for video in video_files:
        video_clip = VideoFileClip(video, audio=False)  # Avoid loading audio for speed optimization

        # Ensure the video has a blurred background using FFmpeg
        video_width, video_height = video_clip.size
        aspect_ratio = video_width / video_height

        if abs(aspect_ratio - 9 / 16) > 0.02 or video_width < 1080 or video_height < 1920:
            # Shorten the temporary file name
            temp_blurred_file = os.path.join(
                os.path.dirname(video), f"blurred_{os.path.basename(video)}"
            )

            # FFmpeg command to blur the video
            ffmpeg_command = [
                "ffmpeg",
                "-i", video,
                "-vf", (
                    "scale=1080:1920,boxblur=luma_radius=20:luma_power=3,"
                    "pad=1080:1920:(ow-iw)/2:(oh-ih)/2:black"
                ),
                "-c:v", "libx264", "-preset", "fast", "-y", temp_blurred_file
            ]

            # Run the FFmpeg command and check for errors
            result = subprocess.run(ffmpeg_command, capture_output=True, text=True)
            if result.returncode != 0:
                print(f"FFmpeg error:\n{result.stderr}")
                raise RuntimeError(f"FFmpeg failed to process the video {video}")

            # Verify the file was created
            if not os.path.exists(temp_blurred_file):
                raise FileNotFoundError(f"Blurred video file not created: {temp_blurred_file}")

            # Load the blurred video file
            video_clip = VideoFileClip(temp_blurred_file, audio=False)

        # If the video is shorter than 3 seconds, take the whole video
        clip_duration = min(video_clip.duration, 3)
        if total_duration + clip_duration > max_duration:
            clip_duration = max_duration - total_duration  # Adjust the last clip to fit the audio length

        clip = video_clip.subclip(0, clip_duration)
        clips.append(clip)
        total_duration += clip_duration

        if total_duration >= max_duration:
            break  # Stop adding clips once the total duration matches the audio length

    # Concatenate the clips and set the audio
    final_video = concatenate_videoclips(clips, method="compose").set_audio(audio_clip)

    # Write the final video to a file
    final_video.write_videofile(output_file, codec="libx264", audio_codec="aac", fps=24, verbose=False, logger=None)

    print(f"Video successfully created: {output_file}")


def create_combined_video(audio_file, output_file, visual_folder="visual", target_resolution=(1080, 1920), transition_ratio=0.1, timestamp=""):
    """Create a single video from timestamped images in the folder with an audio track."""
    update_timestamp()  # Update timestamp on each button click
    if not os.path.exists(audio_file):
        print(f"Error: Audio file {audio_file} not found.")
        return

    # Fetch images containing the timestamp in their filenames
    image_files = [
        os.path.join(visual_folder, f)
        for f in os.listdir(visual_folder)
        if f.lower().endswith(('.png', '.jpg', '.jpeg')) and timestamp in f
    ]

    if not image_files:
        print(f"No images found in folder: {visual_folder} with timestamp {timestamp}")
        return

    # Load the audio file to get its duration
    audio_clip = AudioFileClip(audio_file)
    audio_duration = audio_clip.duration

    # Calculate durations
    num_images = len(image_files)
    image_duration = audio_duration / num_images
    transition_duration = transition_ratio * image_duration

    with tempfile.TemporaryDirectory() as temp_folder:
        # Resize images and prepare clips
        resized_images = [
            resize_image(image_file, target_resolution, temp_folder)
            for image_file in image_files
        ]

        clips = []
        for img in resized_images:
            clip = ImageClip(img).set_duration(image_duration)
            clip = clip.crossfadein(transition_duration).crossfadeout(transition_duration)
            clips.append(clip)

        # Concatenate clips and set audio
        video_clip = concatenate_videoclips(clips, method="compose").set_audio(audio_clip)

        # Ensure final video duration matches audio length
        video_clip = video_clip.set_duration(audio_clip.duration)

        # Write the video to a file
        try:
            video_clip.write_videofile(
                output_file,
                codec="libx264",
                audio_codec="aac",
                fps=24,
                bitrate="10M"  # High-quality output
            )
            print(f"Video successfully created: {output_file}")
            messagebox.showinfo("Success", f"Video saved as {output_file}")
        except Exception as e:
            print(f"Error creating video: {e}")
